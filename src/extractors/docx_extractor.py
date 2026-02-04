"""
DOCX document extractor using python-docx.

Handles Microsoft Word documents (.docx format).
"""

from pathlib import Path

from docx import Document

from ..base_extractor import BaseExtractor
from ..models import TableData, ImageData, DocumentMetadata, FileFormat
from ..utils.markdown_helpers import clean_text, heading_to_markdown, normalize_whitespace


class DOCXExtractor(BaseExtractor):
    """Extracts content from DOCX documents.

    Uses python-docx for all extraction operations.
    """

    def __init__(self, file_path: Path | str) -> None:
        """Initialize DOCX extractor.

        Args:
            file_path: Path to DOCX file.
        """
        super().__init__(file_path)
        self._doc = Document(self.file_path)

    def extract_text(self) -> str:
        """Extract document text as markdown.

        Maps paragraph styles to markdown headings:
        - Title, Heading 1 → # H1
        - Heading 2 → ## H2
        - Heading 3 → ### H3
        - Heading 4 → #### H4
        - Everything else → body text

        Returns:
            Clean markdown string with heading hierarchy preserved.
        """
        lines = []
        for para in self._doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style_name = para.style.name if para.style else ""

            # Map styles to heading levels
            if style_name == "Title" or style_name == "Heading 1":
                lines.append(heading_to_markdown(text, 1))
            elif style_name == "Heading 2":
                lines.append(heading_to_markdown(text, 2))
            elif style_name == "Heading 3":
                lines.append(heading_to_markdown(text, 3))
            elif style_name == "Heading 4":
                lines.append(heading_to_markdown(text, 4))
            else:
                lines.append(text)

        result = "\n\n".join(lines)
        return normalize_whitespace(clean_text(result))

    def extract_tables(self) -> list[TableData]:
        """Extract all tables from the document.

        Returns:
            List of TableData objects with 2D array content.
        """
        tables = []
        for table in self._doc.tables:
            content = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                content.append(row_data)
            if content:
                tables.append(TableData(
                    content=content,
                    page_or_slide=None  # DOCX doesn't expose page numbers
                ))
        return tables

    def extract_images(self) -> list[ImageData]:
        """Extract metadata for all images.

        Accesses images through document relationships.

        Returns:
            List of ImageData objects.
        """
        images = []
        for idx, rel in enumerate(self._doc.part.rels.values()):
            if "image" in rel.reltype:
                try:
                    image_part = rel.target_part
                    # Determine format from content type
                    content_type = image_part.content_type
                    ext = content_type.split("/")[-1]  # e.g., "image/png" -> "png"
                    if ext == "jpeg":
                        ext = "jpg"

                    images.append(ImageData(
                        filename=f"image_{idx}.{ext}",
                        format=ext,
                        width=None,   # Not easily available in python-docx
                        height=None,  # Not easily available in python-docx
                        page_or_slide=None
                    ))
                except Exception:
                    continue  # Skip images that can't be processed
        return images

    def extract_metadata(self) -> DocumentMetadata:
        """Extract document metadata.

        Returns:
            DocumentMetadata object.
        """
        props = self._doc.core_properties
        return DocumentMetadata(
            title=props.title or None,
            author=props.author or None,
            created_date=props.created,   # Already datetime from python-docx
            modified_date=props.modified,  # Already datetime from python-docx
            page_count=None,  # Known limitation: not available in python-docx
            file_format=FileFormat.DOCX,
            file_size_bytes=self.file_path.stat().st_size,
            source_filename=self.file_path.name
        )
